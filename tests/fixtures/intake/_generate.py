"""Generator for tests/fixtures/intake/* binary fixtures.

Ephemeral-toolchain script -- NOT a test-time dependency. Run with:

    uv run --with reportlab --with python-docx python tests/fixtures/intake/_generate.py

Regenerates the committed fixture binaries in this directory. The test suite
(tests/test_intake.py) depends only on the committed binaries and never
invokes this script.
"""

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

FIXTURES_DIR = Path(__file__).resolve().parent


def make_text_layer_pdf(path: Path) -> None:
    """A born-digital PDF with a real, extractable text layer."""
    c = canvas.Canvas(str(path), pagesize=letter)
    c.drawString(72, 700, "Axial intake fixture: this PDF has a real text layer.")
    c.drawString(72, 680, "Second line of body text for good measure.")
    c.showPage()
    c.save()


def make_no_text_layer_pdf(path: Path) -> None:
    """A PDF with NO text layer at all -- vector graphics only.

    No drawString call anywhere, so pypdf's extract_text() returns empty /
    whitespace-only output, simulating a scanned/image-only source without
    needing to embed a raster image.
    """
    c = canvas.Canvas(str(path), pagesize=letter)
    c.rect(72, 600, 200, 100, stroke=1, fill=0)
    c.line(72, 500, 400, 500)
    c.showPage()
    c.save()


def make_text_docx(path: Path) -> None:
    """A DOCX with real body-text paragraphs."""
    doc = Document()
    doc.add_paragraph("Axial intake fixture: this DOCX has real body text.")
    doc.add_paragraph("A second paragraph so there is unambiguous prose content.")
    doc.save(str(path))


def make_unsupported_txt(path: Path) -> None:
    path.write_text("plain text file with an unsupported extension\n", encoding="utf-8")


def make_unsupported_png(path: Path) -> None:
    # Minimal valid 1x1 transparent PNG. Content is irrelevant to the test;
    # only the .png extension matters, but a real PNG avoids ambiguity.
    png_bytes = bytes.fromhex(
        "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c4"
        "890000000a49444154789c6360000002000155a2415e0000000049454e44ae42"
        "6082"
    )
    path.write_bytes(png_bytes)


def main() -> None:
    make_text_layer_pdf(FIXTURES_DIR / "text_layer.pdf")
    make_no_text_layer_pdf(FIXTURES_DIR / "no_text_layer.pdf")
    make_text_docx(FIXTURES_DIR / "text.docx")
    make_unsupported_txt(FIXTURES_DIR / "unsupported.txt")
    make_unsupported_png(FIXTURES_DIR / "unsupported.png")
    print("Generated fixtures in", FIXTURES_DIR)


if __name__ == "__main__":
    main()
