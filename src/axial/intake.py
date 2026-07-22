"""Corpus intake: extension gate + text-layer probe (PRD §5 stage 1, §8 P0-1).

Accepts only `.pdf` and `.docx`. Rejects everything else with a clear,
typed, logged reason. Verifies a real text layer exists before anything
downstream runs -- a scanned/image-only PDF is rejected, never silently
passed through an OCR path (there is none in this slice).

Given an LLM client, intake also runs the holdings-completeness check
(§7.11, §8 P0-1b) over the same text layer, via `axial.holdings`, which
owns the check's own cleaning, prompt and flag shape -- this module's job
is only to build `page_texts`, supply the physical page count, and attach
the resulting flag to `Source`. The check is a model call, so it runs only
for a caller that supplies a client -- and the pipeline supplies one exactly
once per source: `extract()` (every ingest path funnels through it) reads
the persisted record first and passes a client only for a source that has
not been judged yet (`holdings_judged`, issue #303).

Every successful intake also writes the persisted source-metadata record
(§7.12/§7.13, §8 P0-1c/P0-1d) to `data/source_meta/<source_id>.json`,
before any extraction runs: the physical page count, the full sha256 file
hash, the §7.11 holdings flag in full (or an explicit no-flag), and
author/title/date read from the PDF's own embedded metadata and title page
(never the filename). The embedded-metadata half is model-free (pypdf/
python-docx's own document-info readers); the title-page half reuses
`axial.holdings.probe`'s one combined model call (issue #285) rather than a
second pass or a hand-rolled positional heuristic, and that same call
cross-checks the embedded metadata against what the title page actually
says -- a PDF's embedded author/title is sometimes recycled from an
unrelated file, and only something that reads the page can notice.

For a PDF, the record also carries a `publisher` field and a raw
`identifier` (issue #326): an ISBN/DOI captured and checksum-validated from
the front matter (`axial.identifiers`, no network) is looked up against
Open Library/Crossref (`axial.bib_lookup`, cached to disk) and merged into
author/title/date/publisher behind a same-work identity guard that
cross-checks the fetched author against intake's already-known one. Two
things independently withhold the fetch and leave the four fields exactly
as the embedded-metadata/title-page read produced them: the front matter
carrying more than one distinct identifier (ambiguous -- e.g. a multi-
volume work's own series listing several ISBNs; no lookup is even
attempted, since a same-author wrong-volume fetch would pass the guard
below), or the guard rejecting a single, unambiguous fetch whose author
does not overlap intake's already-known one. `identifier` is recorded
either way, for audit.
"""

from __future__ import annotations

import json
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import httpx
from docx import Document
from pypdf import PdfReader

from axial import bib_lookup, identifiers
from axial.holdings import FRONT_MATTER_PAGES
from axial.holdings import probe as _holdings_probe
from axial.llm import LLMClient

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

# Persisted source-metadata record (PRD §7.12, §8 P0-1c): one JSON per
# source in data/source_meta/, keyed by the same deterministic source_id
# the tree/envelope/chunk artifacts use. A plain, cwd-relative path with no
# config-file override, mirroring `axial.extract.TREES_DIR`'s own
# convention (unlike `axial.envelope.ENVELOPES_DIR`, which config can
# redirect -- no caller of this module needs that yet, §7.12 names the path
# directly).
SOURCE_META_DIR = Path("data/source_meta")

# The three §7.13 field states for author/title/date, distinguishable on
# read: a value carries `{"value": ..., "provenance": ...}`; the other two
# states are these literal sentinel strings -- "unavailable" (a read was
# attempted and nothing recoverable was found) versus "not_attempted" (no
# mechanism exists for this field/format combination in this slice, e.g. a
# DOCX's publication date). Never conflated: a consumer checks which of the
# three shapes it got before reading `value`.
UNAVAILABLE = "unavailable"
NOT_ATTEMPTED = "not_attempted"

PROVENANCE_EMBEDDED_METADATA = "embedded metadata"
PROVENANCE_TITLE_PAGE = "title page"

# Record key: has this source's one model-backed intake judgment been made
# and landed? (§7.11 holdings + §7.13 title-page read -- one call answers
# both, so one marker governs both.) It is what makes the check affordable
# in the pipeline (issue #303): `extract()` runs on every pass over every
# source, the judgment is a reasoning-ON call, and the record is where it is
# paid for once and read back thereafter. A `holdings_flag` of `null` cannot
# serve as that marker: it means both "judged complete" and "never judged".
HOLDINGS_CHECKED = "holdings_checked"

# `date` never reads embedded metadata (§7.13): a PDF's own CreationDate/
# ModDate measure when the *file* was produced, not when the *work* was
# published. Its only source is a title-page read -- a model call
# (`axial.holdings.probe`, issue #285) rather than a positional regex: the
# retired deterministic heuristic (a copyright-marker proximity match, plus
# a first-non-blank-line title rule) was measured against the real 30-source
# corpus and found right in ~2 of 13 cases where it fired (#268's own
# pattern); do not rebuild it.


class IntakeError(Exception):
    """Base class for all intake errors."""


class UnsupportedExtensionError(IntakeError):
    """Raised when a file's extension is not among SUPPORTED_EXTENSIONS."""

    def __init__(self, path: Path):
        self.path = path
        self.extension = path.suffix
        super().__init__(
            f"unsupported file extension {self.extension!r} for {path}; "
            f"expected one of {sorted(SUPPORTED_EXTENSIONS)}"
        )


class MissingSourceFileError(IntakeError):
    """Raised when the input path does not exist or is not a file."""

    def __init__(self, path: Path):
        self.path = path
        super().__init__(f"missing or unreadable source file: {path}")


class NoTextLayerError(IntakeError):
    """Raised when a source has no extractable text layer."""

    def __init__(self, path: Path):
        self.path = path
        super().__init__(
            f"no text layer found in {path}; scanned/image-only sources are rejected "
            "(no OCR path in this slice)"
        )


@dataclass
class Source:
    """Source-metadata stub returned on successful intake.

    `holdings_flag` (§7.11, §8 P0-1b) is the holdings-completeness check's
    result: `None` for a source judged complete (and for a caller that
    supplied no client, since the judgment is a model call), otherwise a
    dict recording its measurement -- the source, the concluded document
    kind, the claimed extent and what stated it, the observed page count,
    and the model's stated reason. Never a bare boolean, and never any
    source text (DEC-23). Flag-only: a raised flag never blocks intake or
    alters anything else on this object.
    """

    path: Path
    format: str
    text_layer_ok: bool
    holdings_flag: dict | None = None


def check_extension(path: Path) -> str:
    """Validate `path`'s extension and return the detected format ('pdf'/'docx')."""
    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedExtensionError(path)
    return extension.lstrip(".")


def _pdf_page_texts(path: Path) -> list[str]:
    """One raw text-layer string per physical page of `path`, in reading
    order -- the per-page granularity the holdings-completeness probe needs
    (§7.11) and that a single concatenated string discards."""
    reader = PdfReader(str(path))
    return [page.extract_text() or "" for page in reader.pages]


def _extract_pdf_text(path: Path) -> str:
    return "".join(_pdf_page_texts(path))


def _extract_docx_text(path: Path) -> str:
    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def extract_text_layer(path: Path, fmt: str) -> str:
    """Extract `path`'s raw text layer (`fmt`: 'pdf' or 'docx'). Shared by
    `has_text_layer`'s presence check and by any downstream bounded text
    probe that needs actual text content, not just a boolean -- e.g.
    `axial.drive`'s English-only language-gate probe (issue #239, P0-11c),
    which reuses this rather than reimplementing pdf/docx text extraction."""
    if fmt == "pdf":
        return _extract_pdf_text(path)
    if fmt == "docx":
        return _extract_docx_text(path)
    raise ValueError(f"unknown format {fmt!r}")  # pragma: no cover - guarded by check_extension


def has_text_layer(path: Path, fmt: str) -> bool:
    """Probe `path` (of detected format `fmt`, 'pdf' or 'docx') for real body text."""
    return bool(extract_text_layer(path, fmt).strip())


# =============================================================================
# Source-metadata record (PRD §7.12/§7.13, §8 P0-1c/P0-1d)
# =============================================================================


def source_meta_path(source_id: str, source_meta_dir: Path = SOURCE_META_DIR) -> Path:
    """The write-once-per-intake path for `source_id`'s source-metadata
    record JSON (mirrors `axial.extract.tree_path`/`axial.envelope.envelope_path`)."""
    return source_meta_dir / f"{source_id}.json"


def holdings_judged(source_id: str, source_meta_dir: Path | None = None) -> bool:
    """True when `source_id`'s persisted record already carries the intake
    judgment -- the §7.11 holdings verdict and the §7.13 title-page read that
    the same model call produces.

    This is the once-per-source cache predicate the ingest path reads before
    deciding whether to supply `intake()` a client (§7.12, issue #303). It
    keys off the record having been model-derived at all, not off the flag:
    the flag is `null` both for a source judged complete and for one never
    judged, and the two must not be confused.
    """
    meta_dir = source_meta_dir if source_meta_dir is not None else SOURCE_META_DIR
    try:
        record = json.loads(source_meta_path(source_id, meta_dir).read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return False
    return isinstance(record, dict) and record.get(HOLDINGS_CHECKED) is True


def _clean(value: Any) -> str | None:
    """`value` with internal whitespace collapsed and edges stripped, or
    `None` for an absent/blank/non-string input -- shared normalization for
    both the junk-metadata comparison and the recorded value itself.

    Guards against pypdf handing back a raw `NullObject` (or any other
    non-string sentinel) for a malformed metadata field instead of a plain
    `None` -- issue #285 finding 1: `hall-schroeder-anatomy-of-power.pdf`
    crashed intake outright (`AttributeError: 'NullObject' object has no
    attribute 'split'`) because nothing checked the type before calling
    `.split()`."""
    if not isinstance(value, str):
        return None
    cleaned = " ".join(value.split())
    return cleaned or None


def _plausible_metadata_value(value: Any, *junk: Any) -> str | None:
    """`value`, cleaned, when it is a real bibliographic value -- non-empty
    and not equal (case/whitespace-insensitive) to any of `junk`, the
    document's own producer/creator strings that a PDF writer sometimes
    reuses to auto-fill the author field (§7.13: "a producer string as
    author ... is recorded as unavailable, not passed through"). Returns
    `None` for an empty, whitespace-only, or junk-matching input."""
    cleaned = _clean(value)
    if cleaned is None:
        return None
    normalized = cleaned.casefold()
    for candidate in junk:
        cleaned_candidate = _clean(candidate)
        if cleaned_candidate is not None and normalized == cleaned_candidate.casefold():
            return None
    return cleaned


def _bibliographic_field(value: str | None, provenance: str) -> dict[str, str] | str:
    """The §7.13 field shape for a resolved `value`: `{"value", "provenance"}`
    when a real value was found, else the `UNAVAILABLE` sentinel -- the read
    was attempted (this function is only ever called after an attempt), and
    nothing recoverable was found."""
    return {"value": value, "provenance": provenance} if value else UNAVAILABLE


def _embedded_metadata(fmt: str, path: Path) -> tuple[str | None, str | None]:
    """The file's own embedded author/title, junk-filtered (§7.13) -- the
    model-free half of the bibliographic read. Shared by the title-page
    cross-check prompt (`intake()` states this as the claim to check) and
    `read_bibliographic_fields` (which resolves it into the recorded
    field)."""
    if fmt == "pdf":
        info = PdfReader(str(path)).metadata
        junk = (info.producer if info else None, info.creator if info else None)
        author = _plausible_metadata_value(info.author if info else None, *junk)
        title = _plausible_metadata_value(info.title if info else None, *junk)
        return author, title
    if fmt == "docx":
        props = Document(str(path)).core_properties
        return (
            _plausible_metadata_value(props.author),
            _plausible_metadata_value(props.title),
        )
    raise ValueError(f"unknown format {fmt!r}")  # pragma: no cover - guarded by check_extension


def _resolve_bibliographic_value(
    embedded: str | None, title_page_value: str | None, matches_embedded: bool | None
) -> dict[str, str] | str:
    """One §7.13 field (author or title), resolved from both the file's
    embedded metadata and the model's title-page cross-check (issue #285,
    replacing the deterministic title-page heuristic #268 measured out of
    the design -- it read 2 of 13 real cases correctly):

    - the file carries no embedded value at all -> the title-page reading
      (the model's replacement for the retired first-non-blank-line rule)
      carries the whole answer;
    - the file carries an embedded value and the model read the title page
      and judged it does NOT plausibly name this document ->
      `UNAVAILABLE`, never the embedded value -- a wrong value with
      provenance is worse than an honest blank, and this is the whole point
      of the cross-check (recycled/unrelated embedded metadata, #285
      finding 2);
    - the file carries an embedded value and the model's judgment is
      affirmative or absent (no comparison was made) -> the embedded value
      stands, exactly as before the cross-check existed.
    """
    if embedded is None:
        return _bibliographic_field(title_page_value, PROVENANCE_TITLE_PAGE)
    if matches_embedded is False:
        return UNAVAILABLE
    return _bibliographic_field(embedded, PROVENANCE_EMBEDDED_METADATA)


def read_bibliographic_fields(
    fmt: str,
    embedded_author: str | None,
    embedded_title: str | None,
    *,
    title_page: dict[str, Any] | None = None,
) -> dict[str, dict[str, str] | str]:
    """Resolve author/title/date into their §7.13 record shape from
    `embedded_author`/`embedded_title` (the file's own embedded metadata,
    already junk-filtered by `_embedded_metadata`) and, for a PDF, the model's
    title-page reading (`title_page`, `axial.holdings.probe`'s own shape --
    `None` when no client ran the check this call). Never the filename.

    A DOCX exposes no title-page equivalent to read and no valid `date`
    provenance at all in this slice (its own `created`/`modified` properties
    are file timestamps, not a publication date, exactly like a PDF's
    CreationDate) -- its `date` is `NOT_ATTEMPTED` rather than a guessed
    `UNAVAILABLE`, naming plainly that no mechanism was even tried.
    """
    if fmt == "docx":
        return {
            "author": _bibliographic_field(embedded_author, PROVENANCE_EMBEDDED_METADATA),
            "title": _bibliographic_field(embedded_title, PROVENANCE_EMBEDDED_METADATA),
            "date": NOT_ATTEMPTED,
        }

    if fmt != "pdf":
        raise ValueError(f"unknown format {fmt!r}")  # pragma: no cover - guarded by check_extension

    if title_page is None:
        # No client this call: the title-page read/cross-check is a model
        # call and runs only for a caller that supplies one (mirrors
        # holdings_flag, §7.11/§7.12). Embedded metadata is model-free and
        # trusted as-is when it is all that ran.
        return {
            "author": _bibliographic_field(embedded_author, PROVENANCE_EMBEDDED_METADATA),
            "title": _bibliographic_field(embedded_title, PROVENANCE_EMBEDDED_METADATA),
            "date": UNAVAILABLE,
        }

    return {
        "author": _resolve_bibliographic_value(
            embedded_author, title_page.get("author"), title_page.get("author_matches_embedded")
        ),
        "title": _resolve_bibliographic_value(
            embedded_title, title_page.get("title"), title_page.get("title_matches_embedded")
        ),
        "date": _bibliographic_field(title_page.get("date"), PROVENANCE_TITLE_PAGE),
    }


def _resolve_recorded_field(
    key: str, computed_value: Any, client: LLMClient | None, meta_path: Path
) -> Any:
    """The value to WRITE for record key `key` on this call.

    A caller that supplies no `client` never re-runs a call's model-backed
    judgment (§7.11/§7.13 are paid model calls): `computed_value` reflects
    only what a client-less call can determine on its own. Writing that
    straight through would silently regress an already-recorded, better
    answer the next time any client-less caller re-validates the same
    source -- and `extract()` does exactly that on every call, including
    from inside envelope regeneration -- which is precisely the loss §7.12
    rules out ("does not lose ... the holdings flag"). So a client-less call
    instead PRESERVES whatever `key` the existing on-disk record already
    carries (`computed_value` if there is no existing record, or the key is
    absent from it). A call that DID supply a client always writes its own
    freshly computed answer -- that is a real decision, not an absence of
    one, and must overwrite a stale answer from an earlier, now-superseded
    call."""
    if client is not None:
        return computed_value
    if not meta_path.exists():
        return computed_value
    try:
        existing = json.loads(meta_path.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return computed_value
    if not isinstance(existing, dict) or key not in existing:
        return computed_value
    return existing[key]


def _resolve_holdings_flag(
    computed_flag: dict | None, client: LLMClient | None, meta_path: Path
) -> dict | None:
    """The `holdings_flag` value to WRITE into the record on this call --
    distinct from `Source.holdings_flag`, the per-call return value, which
    stays exactly `computed_flag` (unchanged contract, see `intake`'s own
    docstring). A thin, holdings-flag-typed wrapper over
    `_resolve_recorded_field`, kept as its own name since `intake()`'s and
    this module's own docstrings already refer to it by name."""
    return _resolve_recorded_field("holdings_flag", computed_flag, client, meta_path)


def _resolve_bibliographic_fields(
    computed: dict[str, Any], client: LLMClient | None, meta_path: Path
) -> dict[str, Any]:
    """`author`/`title`/`date`/`publisher` resolved the same way as
    `holdings_flag` (`_resolve_recorded_field`, one call per field): a
    client-less call (every `extract()` validation call, §7.12) preserves
    whatever the existing on-disk record already carries for each field
    rather than regressing it back to a client-less-only read -- the
    title-page cross-check, and the identifier-lookup merge below, that
    produced a better answer both depend on a fuller read (a model call, or
    the freshly-computed same-work identity guard) than a client-less call
    can redo, and a later call that skips it must never silently undo it."""
    return {
        field: _resolve_recorded_field(field, computed[field], client, meta_path)
        for field in ("author", "title", "date", "publisher")
    }


# =============================================================================
# Identifier capture + lookup merge (§7.12/§7.13, issue #326): an ISBN/DOI
# found in a source's own front matter, resolved against Open Library/
# Crossref, and merged into the bibliographic fields.
#
# Two independent safeguards, each catching a DIFFERENT failure mode --
# post-review correction, the founder's own ruling on #326:
# - **Ambiguity abstention** (`identifiers.capture`'s `abstained` shape,
#   applied below before any lookup runs) catches a multi-volume work whose
#   front matter lists several of its own series' ISBNs in one block. This
#   is the one REAL near-miss measured on the corpus
#   (`mann-sources-of-social-power-v1`/`v3`/`v4` all carry the identical
#   ISBN `9781107028654`) -- an author-overlap guard cannot catch it,
#   because the wrong volume shares the true author (the fetch resolves to
#   "Mann, Michael", which overlaps every volume's own known author, so the
#   guard alone would pass it through). Abstaining before the lookup is
#   what protects this case.
# - The **author-overlap guard** (`authors_plausibly_overlap`, kept as one
#   check, no general fuzzy-matching framework) catches a DIFFERENT
#   mismatch: a single, unambiguous identifier that resolves to an entirely
#   different person's work (a mistyped/recycled identifier). It does not
#   and structurally cannot separate same-author volumes from each other --
#   that is ambiguity abstention's job, not this guard's.
# =============================================================================


def _strip_diacritics(text: str) -> str:
    decomposed = unicodedata.normalize("NFKD", text)
    return "".join(ch for ch in decomposed if not unicodedata.combining(ch))


def _name_tokens(name: str) -> set[str]:
    """Meaningful name tokens for the identity guard: diacritics folded away
    (`Siniša` -> `sinisa`), commas normalized so `"Last, First"` and
    `"First Last"` yield the same token set, bare initials/punctuation
    dropped."""
    ascii_name = _strip_diacritics(name).replace(",", " ")
    return {token.casefold() for token in re.findall(r"[A-Za-z]+", ascii_name) if len(token) > 1}


def authors_plausibly_overlap(known_author: str | None, fetched_author: str | None) -> bool:
    """The same-work identity guard: `True` when `known_author` (intake's
    own already-known reading -- embedded metadata or the title-page read)
    and `fetched_author` (the identifier lookup's answer) share at least one
    meaningful name token.

    `known_author=None` (intake has no author of its own to cross-check
    against) passes -- there is nothing to contradict, and rejecting a fetch
    for that reason alone would defeat the real gap this slice fixes
    (`ayubi-over-stating-the-arab-state`'s `None` title, where no prior
    author reading existed either). A present `known_author` with no
    overlapping token in `fetched_author` fails -- catching a single,
    unambiguous identifier that resolves to an entirely different person's
    work.

    This guard does **not** catch a same-author wrong-volume/wrong-edition
    mismatch (measured: a fetch for the wrong Mann volume still resolves to
    "Mann, Michael", which overlaps the known "Michael Mann" -- this
    function returns `True`, correctly, for that pair). That failure mode is
    caught upstream instead, by ambiguity abstention on the capture itself
    (see the section docstring above) -- not by this function."""
    if not known_author:
        return True
    if not fetched_author:
        return False
    return bool(_name_tokens(known_author) & _name_tokens(fetched_author))


def _known_author_value(author_field: dict[str, str] | str) -> str | None:
    return author_field.get("value") if isinstance(author_field, dict) else None


def _capture_identifier(fmt: str, page_texts: list[str]) -> dict[str, Any] | None:
    """The validated ISBN/DOI `identifier` this source's front matter
    carries (see `identifiers.capture` for its three possible shapes,
    including the `abstained` ambiguous-capture case), or `None` -- PDF
    only (issue #326): identifier scanning reads the same head window
    `axial.holdings` already uses for its own front-matter judgment
    (`FRONT_MATTER_PAGES`), adding no new tunable window size of its own."""
    if fmt != "pdf":
        return None
    return identifiers.capture("\n".join(page_texts[:FRONT_MATTER_PAGES]))


def _lookup_identifier(
    identifier: dict[str, Any],
    bib_transport: httpx.BaseTransport | None,
    bib_cache_dir: Path | None,
) -> dict[str, Any]:
    if identifier["type"] == "isbn":
        return bib_lookup.resolve_isbn(
            identifier["value"], transport=bib_transport, cache_dir=bib_cache_dir
        )
    return bib_lookup.resolve_doi(
        identifier["value"], transport=bib_transport, cache_dir=bib_cache_dir
    )


def _merge_identifier_fields(
    biblio: dict[str, Any],
    identifier: dict[str, Any] | None,
    fmt: str,
    bib_transport: httpx.BaseTransport | None,
    bib_cache_dir: Path | None,
) -> dict[str, Any]:
    """`biblio` (already resolved from embedded metadata/title page), with
    the identifier-lookup override applied when the same-work identity
    guard passes. `publisher` has no reading mechanism other than this
    lookup: `NOT_ATTEMPTED` for a DOCX (no identifier scan runs at all in
    this slice) and `UNAVAILABLE` for a PDF with no identifier, an
    ambiguous capture, an unresolved lookup, or a failing guard (a read was
    attempted -- the scan, or the scan plus the lookup -- and nothing
    recoverable/trustworthy came of it).

    An **ambiguous** `identifier` (`identifiers.capture`'s `abstained: True`
    shape -- more than one distinct identifier found, e.g. a multi-volume
    work's "also available" ISBN block) never reaches the lookup at all:
    no single identifier exists to look up, and the same-work identity
    guard cannot rescue this case downstream -- a wrong-volume fetch
    typically shares the true author (measured: `mann-sources-of-social-
    power-v1`/`v3`/`v4` all carry ISBN `9781107028654`, which resolves to
    author "Mann, Michael", which overlaps every volume's own known author,
    so the guard alone would pass a wrong-volume fetch through). Abstaining
    on the ambiguous capture, before any lookup runs, is what protects this
    case -- not the guard."""
    merged = dict(biblio)
    merged["publisher"] = NOT_ATTEMPTED if fmt != "pdf" else UNAVAILABLE
    if identifier is None or identifier.get("abstained"):
        return merged

    lookup = _lookup_identifier(identifier, bib_transport, bib_cache_dir)
    if not lookup.get("resolved"):
        return merged

    known_author = _known_author_value(biblio["author"])
    if not authors_plausibly_overlap(known_author, lookup.get("author")):
        return merged

    provenance = lookup["source"]
    merged["author"] = _bibliographic_field(lookup.get("author"), provenance)
    merged["title"] = _bibliographic_field(lookup.get("title"), provenance)
    merged["date"] = _bibliographic_field(lookup.get("date"), provenance)
    merged["publisher"] = _bibliographic_field(lookup.get("publisher"), provenance)
    return merged


def build_source_meta(
    source_id: str,
    path: Path,
    fmt: str,
    physical_pages: int | None,
    holdings_flag: dict | None,
    bibliographic_fields: dict[str, Any],
    holdings_checked: bool,
    identifier: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Assemble the §7.12 source-metadata record: the full sha256 file hash,
    the physical page count (present for a PDF, an explicit `null` for a
    DOCX -- distinct from a numeric zero), the §7.11 holdings flag in full
    or an explicit no-flag, whether that judgment has been made at all
    (`holdings_checked`, see `HOLDINGS_CHECKED`), the already-resolved §7.13
    bibliographic fields (author/title/date/publisher), and the raw
    `identifier` found in the source's own front matter (issue #326) --
    kept for audit even when it was not used for the four fields above.
    `identifier` is one of: `None` (nothing found); `{"type", "value"}` (one
    identifier found, the same-work identity guard may have accepted or
    rejected it); or `{"type", "value": None, "abstained": True,
    "candidates": [...]}` (more than one distinct identifier found --
    ambiguous, no lookup was even attempted, see `identifiers.capture`).
    Carries no source text (DEC-23): values and short reasons only."""
    from axial.envelope import content_digest  # local import: dodges the
    # envelope -> extract -> intake import cycle (mirrors extract.py's own
    # local `compute_source_id` import, same reason).

    record: dict[str, Any] = {
        "source_id": source_id,
        "file_hash": content_digest(path),
        "physical_page_count": physical_pages,
        "holdings_flag": holdings_flag,
        HOLDINGS_CHECKED: holdings_checked,
        "identifier": identifier,
    }
    record.update(bibliographic_fields)
    return record


def write_source_meta(record: dict[str, Any], path: Path) -> None:
    """Write the source-metadata record JSON, creating parent directories as
    needed (mirrors `axial.envelope.write_envelope`'s own convention)."""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(record, indent=2, sort_keys=True) + "\n", encoding="utf-8")


def intake(
    path: str | Path,
    *,
    client: LLMClient | None = None,
    source_meta_dir: Path | None = None,
    bib_transport: httpx.BaseTransport | None = None,
    bib_cache_dir: Path | None = None,
) -> Source:
    """Run intake on `path`: validate extension, verify a text layer, and --
    when `client` is given -- run the holdings-completeness check
    (§7.11/§8 P0-1b) and attach its flag. A raised flag is flag-only: it
    never raises, never rejects, and the source still completes intake
    exactly as an unflagged one would.

    A DOCX is checked too (the earlier blanket DOCX exemption is retired,
    §7.11); it exposes no physical page count, which the check handles as
    unobtainable evidence rather than as damning.

    For a PDF, an ISBN/DOI found in the front matter is also looked up
    against Open Library/Crossref (issue #326) and merged into
    author/title/date/publisher (`_merge_identifier_fields`) -- independent
    of `client`, since the lookup is a free, cached network call, not a
    paid model call. An ambiguous capture (more than one distinct
    identifier found) abstains before any lookup runs; a single,
    unambiguous fetch is still gated by the same-work identity guard.
    `bib_transport` is a mockable `httpx` transport for that lookup
    (`httpx.MockTransport`, the same seam `axial.llm.OpenRouterClient`
    already uses); `None` uses a real `httpx.Client` in production.
    `bib_cache_dir` overrides where the raw lookup response is cached
    (defaults to `axial.bib_lookup.CACHE_DIR`), mirroring `source_meta_dir`
    -- so a test can keep every write hermetic under its own `tmp_path`
    rather than the real, gitignored `data/bib_lookup_cache/`.

    Every successful call also writes the persisted source-metadata record
    (§7.12/§7.13) to `<source_meta_dir>/<source_id>.json` -- `source_meta_dir`
    defaults to `SOURCE_META_DIR` -- before returning, so the record exists
    before any extraction runs. This is unconditional and does not depend on
    `client`: the page count, file hash, and bibliographic fields are
    model-free and always read; only `holdings_flag` depends on whether this
    call ran the judgment (see `_resolve_holdings_flag`). The record also
    carries whether that judgment has landed at all (`HOLDINGS_CHECKED`), so
    the ingest path can pay for it once per source and read it back after
    (`holdings_judged`, issue #303)."""
    path = Path(path)

    if not path.is_file():
        raise MissingSourceFileError(path)

    fmt = check_extension(path)

    if fmt == "pdf":
        page_texts = _pdf_page_texts(path)
        if not "".join(page_texts).strip():
            raise NoTextLayerError(path)
        physical_pages: int | None = len(page_texts)
    else:
        text = extract_text_layer(path, fmt)
        if not text.strip():
            raise NoTextLayerError(path)
        page_texts = [text]
        physical_pages = None

    embedded_author, embedded_title = _embedded_metadata(fmt, path)

    if client is None:
        holdings_flag: dict | None = None
        title_page: dict[str, Any] | None = None
        answered = False
    else:
        probed = _holdings_probe(
            page_texts,
            client=client,
            physical_pages=physical_pages,
            source_name=path.name,
            # A DOCX has no title page to cross-check; only a PDF's
            # embedded claim is worth stating in the prompt (§7.13).
            embedded_author=embedded_author if fmt == "pdf" else None,
            embedded_title=embedded_title if fmt == "pdf" else None,
        )
        holdings_flag = probed["holdings_flag"]
        title_page = probed["title_page"] if fmt == "pdf" else None
        answered = bool(probed["answered"])

    # §7.12/§7.13, §8 P0-1c/P0-1d: local import dodges the
    # envelope -> extract -> intake import cycle (see build_source_meta's
    # own docstring for the identical reasoning extract.py already follows).
    from axial.envelope import compute_source_id

    source_id = compute_source_id(path)
    meta_dir = source_meta_dir if source_meta_dir is not None else SOURCE_META_DIR
    meta_path = source_meta_path(source_id, meta_dir)
    written_flag = _resolve_holdings_flag(holdings_flag, client, meta_path)
    written_checked = _resolve_recorded_field(HOLDINGS_CHECKED, answered, client, meta_path)
    biblio = read_bibliographic_fields(fmt, embedded_author, embedded_title, title_page=title_page)
    identifier = _capture_identifier(fmt, page_texts)
    biblio = _merge_identifier_fields(biblio, identifier, fmt, bib_transport, bib_cache_dir)
    written_biblio = _resolve_bibliographic_fields(biblio, client, meta_path)
    record = build_source_meta(
        source_id,
        path,
        fmt,
        physical_pages,
        written_flag,
        written_biblio,
        bool(written_checked),
        identifier,
    )
    write_source_meta(record, meta_path)

    return Source(path=path, format=fmt, text_layer_ok=True, holdings_flag=holdings_flag)
