"""Export the finished ask as one downloadable file (issue #724): the
reader-facing answer (`axial.answer.reader.render_reader_answer`) as
markdown -- and, converted from that SAME markdown, `docx` and `odt`.

**The metrics left the document in #783.** This module used to export
`render_markdown(record) + _render_metrics_markdown(metrics)`: the AUDIT
rendering plus a telemetry appendix, in the file a reader downloads. So an
exported `.docx` carried lines like `evidence_share=0.523 available_share=
0.318 usage_ratio=22.954545454545453`, which is the exact failure
`axial.paper.render`'s own docstring names. The metrics are still computed,
still served beside the record by `GET /asks/{id}/paper`
(`metrics_block`, below) and still rendered in the web client's metrics
panel; they are no longer stapled to the document.

**One rendering path, three containers, per the issue's own line.**
`render_export_markdown` is the only place record content turns into
text; `render_docx`/`render_odt` each walk that text's own bounded set of
constructs (`#`/`##` headings, `- ` bullets, plain paragraphs, `**bold**`
spans -- exactly what `render_markdown` and this module's own metrics
appendix ever emit, nothing else) into their own binary format. Neither
converter re-reads the record: a citation mode's redaction (issue #690,
already applied by the caller before this module ever runs, on `record`)
can never diverge between the three containers, because all three are
built from the one already-redacted string.

**Conversion library, decided here.** `docx` uses `python-docx` (already
a main dependency, for the extract path's own `.docx` reading). `odt`
uses `odfpy` (added for this issue, pure Python, no binary). Pandoc would
also produce both, at the cost of a binary #691's image would have to
carry and `.env.example` would have to name -- and this module's own
five-construct line walk covers everything either format here ever needs
to render, so the extra binary buys nothing this issue's "done when"s
need. `pypandoc` was therefore not the right call.

**Export is free** (the issue's own line): nothing here calls a model,
touches a job row's `cost_usd`/quota, or opens a database connection --
`axial.service.api`'s export route reads the already-persisted record and
never charges or counts it.
"""

from __future__ import annotations

import io
import re
from typing import Any

from axial.answer.reader import render_reader_answer

EXPORT_FORMATS = ("md", "docx", "odt")

_METRICS_FIELDS = ("cost", "model_by_pass", "coverage_map", "confidence")

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def metrics_block(record: dict[str, Any]) -> dict[str, Any]:
    """The four §7.3 fields `GET /asks/{id}/paper` and this module both
    serve beside the record (issue #724's own scope line): `cost`,
    `model_by_pass`, `coverage_map`, `confidence` -- never `retries` or a
    shape band, which are Phase-C-only (`src/axial/paper/record.py`) and do
    not exist on an ask's own §7.3 record at all."""
    return {field: record.get(field) for field in _METRICS_FIELDS}


def _strip_markdown_emphasis(text: str) -> str:
    """`**bold**` markers removed, plain text kept -- neither `docx` nor
    `odt` needs the emphasis styling to meet this issue's bar, and
    stripping is one regex rather than a second rendering path that tracks
    bold runs through both converters."""
    return _BOLD_RE.sub(r"\1", text)


def render_export_markdown(record: dict[str, Any]) -> str:
    """The exported document: the reader-facing rendering of `record`
    (`axial.answer.reader.render_reader_answer`) and nothing else. `record`
    should already be rendered through this deployment's citation mode
    (`axial.service.citation.render_record_for_serving`) by the caller,
    exactly as `GET /asks/{id}/paper` already does -- and since #783 that
    now MATTERS to what comes out: the reader render prints each claim's
    citation from the resolved `citation` block, so an unresolved record
    exports raw `chunk:<id>` pointers rather than `Vignal 2021, ch. 30`.
    Reads only `record`, calls no model, opens no vault -- the same purity
    `render_reader_answer` itself holds."""
    return render_reader_answer(record)


def render_docx(markdown_text: str) -> bytes:
    """`markdown_text` (this module's own bounded output) as a `.docx`
    file's bytes, via `python-docx` (module docstring)."""
    from docx import Document

    document = Document()
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("## "):
            document.add_heading(_strip_markdown_emphasis(stripped[3:]), level=2)
        elif stripped.startswith("# "):
            document.add_heading(_strip_markdown_emphasis(stripped[2:]), level=1)
        elif stripped.startswith("- "):
            document.add_paragraph(_strip_markdown_emphasis(stripped[2:]), style="List Bullet")
        else:
            document.add_paragraph(_strip_markdown_emphasis(stripped))
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_odt(markdown_text: str) -> bytes:
    """The same walk as `render_docx`, into `.odt` via `odfpy` -- pure
    Python, no binary (module docstring)."""
    from odf.opendocument import OpenDocumentText
    from odf.text import H, List, ListItem, P

    document = OpenDocumentText()
    bullet_list = None
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            bullet_list = None
            continue
        if stripped.startswith("## "):
            document.text.addElement(H(outlinelevel=2, text=_strip_markdown_emphasis(stripped[3:])))
            bullet_list = None
        elif stripped.startswith("# "):
            document.text.addElement(H(outlinelevel=1, text=_strip_markdown_emphasis(stripped[2:])))
            bullet_list = None
        elif stripped.startswith("- "):
            if bullet_list is None:
                bullet_list = List()
                document.text.addElement(bullet_list)
            item = ListItem()
            item.addElement(P(text=_strip_markdown_emphasis(stripped[2:])))
            bullet_list.addElement(item)
        else:
            document.text.addElement(P(text=_strip_markdown_emphasis(stripped)))
            bullet_list = None
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
