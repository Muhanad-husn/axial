"""Convert the generated report markdown into a .docx.

Handles exactly the subset of markdown the generator emits: headings, pipe
tables, bullets with one nesting level, blockquotes, images, horizontal rules,
and inline bold / italic / code / links.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import sys

REPORTS = Path(__file__).resolve().parents[1]

INK = RGBColor(0x1F, 0x29, 0x33)
ACCENT = RGBColor(0xC8, 0x55, 0x3D)

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))")


def add_runs(par, text: str, base_italic: bool = False, base_bold: bool = False):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            # Recurse so a link or code span nested inside **bold** still
            # renders as a link or code span, just also bold.
            add_runs(par, part[2:-2], base_italic=base_italic, base_bold=True)
            continue
        if part.startswith("*") and part.endswith("*") and len(part) > 2:
            # Same, for *italic* — this is what a link nested inside an
            # italic lead-in paragraph needs to still render as a link.
            add_runs(par, part[1:-1], base_italic=True, base_bold=base_bold)
            continue
        if part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        elif part.startswith("["):
            m = re.match(r"\[([^\]]+?)\]\(([^)]+?)\)", part)
            r = par.add_run(m.group(1))
            r.font.color.rgb = ACCENT
        else:
            r = par.add_run(part)
        if base_italic:
            r.italic = True
        if base_bold:
            r.bold = True


def style_document(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15
    for i, size in [(1, 22), (2, 16), (3, 13), (4, 11)]:
        st = doc.styles[f"Heading {i}"]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = INK if i > 1 else ACCENT
        st.font.bold = True
        st.paragraph_format.space_before = Pt(16 if i < 3 else 12)
        st.paragraph_format.space_after = Pt(6)


def flush_table(doc, rows):
    if len(rows) < 2:
        return
    header, body = rows[0], rows[2:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for cell, text in zip(table.rows[0].cells, header):
        cell.text = ""
        add_runs(cell.paragraphs[0], text)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for line in body:
        cells = table.add_row().cells
        for cell, text in zip(cells, line):
            cell.text = ""
            add_runs(cell.paragraphs[0], text)
    for row in table.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                par.paragraph_format.space_after = Pt(2)
                for r in par.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def split_row(line: str):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(src: Path, dst: Path):
    text = src.read_text(encoding="utf-8")
    doc = Document()
    style_document(doc)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    sec.left_margin = sec.right_margin = Inches(0.95)

    lines = text.splitlines()
    table_buf: list[list[str]] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            # A fenced block that survived substitution (e.g. inline code
            # samples). Render as a monospace paragraph rather than letting
            # each line fall through as its own prose paragraph.
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing fence
            if code_lines:
                par = doc.add_paragraph()
                r = par.add_run("\n".join(code_lines))
                r.font.name = "Consolas"
                r.font.size = Pt(9)
            continue

        if stripped.startswith("|"):
            table_buf.append(split_row(stripped))
            i += 1
            continue
        if table_buf:
            flush_table(doc, table_buf)
            table_buf = []

        if not stripped:
            i += 1
            continue

        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if m:
            # Resolved against REPORTS, not src.parent: a spliced build copy
            # under build/ still references images by their path from REPORTS.
            img = REPORTS / m.group(2)
            if img.exists():
                # Diagrams get the full content width; the logo stays small.
                width = Inches(6.35) if "diagrams/" in m.group(2) else Inches(1.1)
                par = doc.add_paragraph()
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                par.add_run().add_picture(str(img), width=width)
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        m = re.match(r"(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            par = doc.add_paragraph(style=f"Heading {min(level, 4)}")
            add_runs(par, m.group(2))
            i += 1
            continue

        if stripped.startswith("> "):
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.3)
            par.paragraph_format.space_before = Pt(4)
            add_runs(par, stripped[2:], base_italic=True)
            i += 1
            continue

        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            par = doc.add_paragraph(style=style)
            par.paragraph_format.space_after = Pt(3)
            add_runs(par, m.group(2))
            i += 1
            continue

        par = doc.add_paragraph()
        add_runs(par, stripped)
        i += 1

    if table_buf:
        flush_table(doc, table_buf)

    doc.save(dst)
    print("wrote", dst)


def main():
    stem = sys.argv[1] if len(sys.argv) > 1 else "axial-report"
    convert(REPORTS / f"{stem}.md", REPORTS / f"{stem}.docx")


if __name__ == "__main__":
    main()
